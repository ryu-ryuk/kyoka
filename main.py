import aiohttp
import asyncio
from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_pinecone import PineconeVectorStore
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever
from langchain.docstore.document import Document
from langchain_huggingface import HuggingFaceEmbeddings
import fitz
import tiktoken
from pinecone import Pinecone as PineconeClient, ServerlessSpec
import logging
from typing import List, Dict, Any
import time
from flashrank import Ranker, RerankRequest
import torch
import os
import hashlib
import json
import google.generativeai as genai
import re
from io import BytesIO

# doc parse
from docx import Document as DocxDocument
import email
from email.policy import default

from dotenv import load_dotenv

# conosnnsntantsts
from constants import *

load_dotenv()

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

app = FastAPI(title="hackrx query system", description="fast llm document query system")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# global caches
DOCUMENT_CACHE = {}
RETRIEVER_CACHE = {}


class QueryRequest(BaseModel):
    documents: str
    questions: List[str]


class Answer(BaseModel):
    query: str
    answer: str
    rationale: str
    source: str


class QueryResponse(BaseModel):
    answers: List[Answer]


class SimpleQueryResponse(BaseModel):
    answers: List[str]


def initialize_components():
    global embeddings, tokenizer, pinecone_client, index, ranker, gemini_model

    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_NAME,
        model_kwargs={"device": "cuda" if torch.cuda.is_available() else "cpu"},
        encode_kwargs={"batch_size": ENCODING_BATCH_SIZE},
    )

    tokenizer = tiktoken.get_encoding("cl100k_base")
    pinecone_client = PineconeClient(api_key=PINECONE_API_KEY)
    index = initialize_pinecone_index(pinecone_client)

    try:
        ranker = Ranker(model_name=RERANKER_MODEL_NAME, max_length=512)
    except:
        ranker = None

    if GEMINI_API_KEY:
        gemini_model = genai.GenerativeModel(LLM_MODEL_NAME)
    else:
        gemini_model = None

    logger.info(f"initialized. gpu: {torch.cuda.is_available()}")


def initialize_pinecone_index(pinecone_client: PineconeClient):
    try:
        existing_indexes = pinecone_client.list_indexes().names()

        if PINECONE_INDEX_NAME not in existing_indexes:
            pinecone_client.create_index(
                name=PINECONE_INDEX_NAME,
                dimension=EMBEDDING_DIMENSIONS,
                metric=PINECONE_METRIC,
                spec=ServerlessSpec(cloud=PINECONE_CLOUD, region=PINECONE_REGION),
            )
            time.sleep(INDEX_CREATION_WAIT_TIME)

        return pinecone_client.Index(PINECONE_INDEX_NAME)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"pinecone setup failed: {str(e)}")


def detect_file_type(file_content: bytes, filename: str = None) -> str:
    if file_content.startswith(b"%PDF"):
        return "pdf"
    elif file_content.startswith(b"PK"):
        return "docx"
    elif (
        b"from:" in file_content[:1000].lower() or b"to:" in file_content[:1000].lower()
    ):
        return "email"

    if filename:
        ext = filename.lower().split(".")[-1]
        if ext in ["pdf", "docx", "eml", "txt"]:
            return ext if ext != "txt" else "email"

    return "unknown"


def extract_text_from_docx(file_content: bytes, doc_id: str) -> List[Document]:
    try:
        doc = DocxDocument(BytesIO(file_content))
        text_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # simple heading detection
                if len(text) < HEADING_LENGTH_THRESHOLD and not text.endswith("."):
                    text_parts.append(f"SECTION: {text}")
                else:
                    text_parts.append(text)

        full_text = "\n".join(text_parts)
        chunks = create_semantic_chunks(full_text, doc_id, 1, "document")

        return chunks
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"docx extraction failed: {str(e)}")


def extract_text_from_email(file_content: bytes, doc_id: str) -> List[Document]:
    try:
        email_text = file_content.decode("utf-8", errors="ignore")
        msg = email.message_from_string(email_text, policy=default)

        subject = msg.get("Subject", "no subject")
        sender = msg.get("From", "unknown")

        # extract body
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    body += part.get_payload(decode=True).decode(
                        "utf-8", errors="ignore"
                    )
        else:
            if msg.get_content_type() == "text/plain":
                body = msg.get_payload(decode=True).decode("utf-8", errors="ignore")

        email_content = f"SECTION: email\nsubject: {subject}\nfrom: {sender}\n\n{body}"
        chunks = create_semantic_chunks(email_content, doc_id, 1, "email")

        return chunks
    except Exception as e:
        # fallback to plain text
        text_content = file_content.decode("utf-8", errors="ignore")
        chunks = create_semantic_chunks(text_content, doc_id, 1, "text")
        return chunks


def extract_structured_text_from_pdf(pdf_content: bytes, doc_id: str) -> List[Document]:
    try:
        doc = fitz.open(stream=pdf_content, filetype="pdf")
        documents = []

        for page_num, page in enumerate(doc):
            text_dict = page.get_text("dict")
            page_content = []
            current_section = ""

            for block in text_dict["blocks"]:
                if "lines" in block:
                    block_text = ""
                    font_sizes = []

                    for line in block["lines"]:
                        line_text = ""
                        for span in line["spans"]:
                            text = span["text"].strip()
                            font_sizes.append(span["size"])
                            line_text += text + " "
                        if line_text.strip():
                            block_text += line_text.strip() + "\n"

                    if block_text.strip():
                        avg_font_size = (
                            sum(font_sizes) / len(font_sizes) if font_sizes else 12
                        )
                        is_heading = (
                            avg_font_size > HEADING_FONT_SIZE_THRESHOLD
                            or len(block_text.strip()) < HEADING_LENGTH_THRESHOLD
                        )

                        if (
                            is_heading
                            and len(block_text.strip()) < HEADING_LENGTH_THRESHOLD
                        ):
                            current_section = block_text.strip()
                            page_content.append(f"SECTION: {current_section}")
                        else:
                            page_content.append(block_text.strip())

            full_text = "\n".join(page_content)
            chunks = create_semantic_chunks(
                full_text, doc_id, page_num + 1, current_section
            )
            documents.extend(chunks)

        doc.close()
        return documents
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"pdf extraction failed: {str(e)}")


# splits a long text into smaller, meaningful chunks that:
# - wont exceed the token limit
# - retain semantic context
# - attaches useful metadata for later use

# we use RecursiveCharacterTextSplitter from langchain
#    splits using multiple fallback separators (section > paragraph > sentence > word)
#    respects token length using tokenizer.encode() to count tokens, not characters
#    adds overlap for context continuity between chunks


def create_semantic_chunks(
    text: str, doc_id: str, page_num: int, section: str
) -> List[Document]:
    try:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=MAX_CHUNK_TOKENS,
            chunk_overlap=CHUNK_OVERLAP_TOKENS,
            length_function=lambda x: len(tokenizer.encode(x)),
            separators=["\n\nSECTION:", "\n\n", "\n", ". ", " "],
            keep_separator=True,
        )

        chunks = splitter.split_text(text)
        documents = []

        for i, chunk in enumerate(chunks):
            if chunk.strip():
                clause_matches = re.findall(CLAUSE_PATTERN, chunk.lower())
                clause_ids = list(set(clause_matches)) if clause_matches else []

                metadata = {
                    "doc_id": doc_id,
                    "page": page_num,
                    "chunk_id": f"{page_num}_{i}",
                    "section": section,
                    "clause_ids": clause_ids,
                    "token_count": len(tokenizer.encode(chunk)),
                }

                documents.append(
                    Document(page_content=chunk.strip(), metadata=metadata)
                )

        return documents
    except:
        return []


def extract_structured_text_from_document(
    file_content: bytes, doc_id: str, filename: str = None
) -> List[Document]:
    file_type = detect_file_type(file_content, filename)

    if file_type == "pdf":
        return extract_structured_text_from_pdf(file_content, doc_id)
    elif file_type == "docx":
        return extract_text_from_docx(file_content, doc_id)
    elif file_type == "email":
        return extract_text_from_email(file_content, doc_id)
    else:
        raise HTTPException(status_code=400, detail=ERROR_MESSAGES["unsupported_file"])


async def setup_hybrid_retriever_cached(documents: List[Document], doc_id: str):
    if doc_id in RETRIEVER_CACHE:
        return RETRIEVER_CACHE[doc_id]

    try:
        vectorstore = PineconeVectorStore(
            index=index, embedding=embeddings, text_key="text"
        )

        # faster batch processing
        for i in range(0, len(documents), BATCH_SIZE):
            batch = documents[i : i + BATCH_SIZE]
            try:
                vectorstore.add_documents(batch)
                await asyncio.sleep(API_RATE_LIMIT_DELAY)
            except Exception as batch_error:
                continue

        vector_retriever = vectorstore.as_retriever(
            search_kwargs={"k": VECTOR_RETRIEVAL_K, "include_metadata": True}
        )
        bm25_retriever = BM25Retriever.from_documents(documents, k=BM25_RETRIEVAL_K)

        ensemble_retriever = EnsembleRetriever(
            retrievers=[vector_retriever, bm25_retriever],
            weights=ENSEMBLE_WEIGHTS,
        )

        RETRIEVER_CACHE[doc_id] = ensemble_retriever
        return ensemble_retriever

    except Exception as e:
        # fast fallback
        bm25_retriever = BM25Retriever.from_documents(documents, k=BM25_RETRIEVAL_K)
        RETRIEVER_CACHE[doc_id] = bm25_retriever
        return bm25_retriever


async def get_relevant_context_with_reranking(query: str, retriever) -> List[Document]:
    try:
        chunks = await asyncio.get_event_loop().run_in_executor(
            None, retriever.invoke, query
        )

        if ranker and len(chunks) > MAX_CONTEXT_CHUNKS:
            try:
                passages = [
                    {"id": i, "text": chunk.page_content}
                    for i, chunk in enumerate(chunks[:MAX_RERANK_CANDIDATES])
                ]
                request = RerankRequest(query=query, passages=passages)
                ranked = ranker.rerank(request)

                return [chunks[r["id"]] for r in ranked[:MAX_CONTEXT_CHUNKS]]
            except:
                pass

        return chunks[:MAX_CONTEXT_CHUNKS]
    except:
        return []


async def generate_structured_answer(
    query: str, chunks: List[Document]
) -> Dict[str, str]:
    if not gemini_model or not chunks:
        return {
            "answer": ERROR_MESSAGES["gemini_not_configured"],
            "rationale": "api not configured",
            "source": "n/a",
        }

    context_parts = []
    sources = set()

    for chunk in chunks:
        metadata = chunk.metadata if hasattr(chunk, "metadata") else {}
        page = metadata.get("page", "n/a")
        sources.add(f"page {page}")
        context_parts.append(f"[page {page}]\n{chunk.page_content}")

    context = "\n\n".join(context_parts)
    source_summary = ", ".join(sorted(sources))

    # have to prevent nested json
    prompt = f"""You are analyzing an insurance policy document. Answer the question using only the provided context.

CONTEXT:
{context}

QUESTION: {query}

INSTRUCTIONS:
- Provide a direct, factual answer with specific details
- Include time periods, amounts, conditions, and percentages when mentioned
- Base your answer only on the provided context
- Keep the response concise but complete

Return your response as a valid JSON object with exactly these three fields:
- answer: your factual response to the question
- rationale: brief explanation of your reasoning
- source: page references from the context

JSON Response:"""

    try:
        response = await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: gemini_model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    max_output_tokens=GEMINI_MAX_OUTPUT_TOKENS,
                    temperature=GEMINI_TEMPERATURE,
                ),
            ),
        )

        response_text = response.text.strip()

        # aggressive cleanup for gemini's formatting quirks
        if response_text.startswith("```"):
            # remove markdown code block wrappers
            lines = response_text.split("\n")
            start_idx = 1 if lines[0].strip().startswith("```") else 0
            end_idx = len(lines) - 1 if lines[-1].strip() == "```" else len(lines)
            response_text = "\n".join(lines[start_idx:end_idx]).strip()

        # remove any leading lines like "JSON Response:" or explanation text
        if not response_text.startswith("{"):
            lines = response_text.split("\n")
            for i, line in enumerate(lines):
                if line.strip().startswith("{"):
                    response_text = "\n".join(lines[i:]).strip()
                    break

        try:
            result = json.loads(response_text)

            # fix nested json strings in answer field
            answer = result.get("answer", "")
            if isinstance(answer, str) and answer.strip().startswith("{"):
                try:
                    nested = json.loads(answer)
                    if isinstance(nested, dict) and "answer" in nested:
                        result["answer"] = nested["answer"]
                        result["rationale"] = nested.get(
                            "rationale", result.get("rationale", "")
                        )
                        result["source"] = nested.get(
                            "source", result.get("source", source_summary)
                        )
                except:
                    pass  # keep original if nested parsing fails

            return {
                "answer": result.get("answer", "parsing error"),
                "rationale": result.get("rationale", "based on policy analysis"),
                "source": result.get("source", source_summary),
            }

        except json.JSONDecodeError as e:
            logger.warning(f"json parse error: {e}")
            # extract answer from malformed response
            clean_text = (
                response_text.replace("\n", " ")
                .replace("{", "")
                .replace("}", "")
                .replace('"', "")
            )
            return {
                "answer": clean_text[:300] + "..."
                if len(clean_text) > 300
                else clean_text,
                "rationale": "based on policy document analysis",
                "source": source_summary,
            }

    except Exception as e:
        logger.error(f"gemini error: {str(e)}")
        return {
            "answer": "error generating response",
            "rationale": f"api error: {str(e)}",
            "source": source_summary,
        }


async def download_pdf_cached(url: str) -> bytes:
    url_hash = hashlib.md5(url.encode()).hexdigest()
    if url_hash in DOCUMENT_CACHE:
        return DOCUMENT_CACHE[url_hash]

    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise HTTPException(
                        status_code=400, detail=ERROR_MESSAGES["download_failed"]
                    )

                file_content = await response.read()
                DOCUMENT_CACHE[url_hash] = file_content
                return file_content
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"download failed: {str(e)}")


@app.post("/hackrx/run", response_model=QueryResponse)
async def run_query_url(request: QueryRequest):
    return await process_queries(request.documents, request.questions, is_url=True)


@app.post("/hackrx/upload", response_model=QueryResponse)
async def run_query_upload(questions: List[str] = [], file: UploadFile = File(...)):
    file_type = detect_file_type(await file.read(), file.filename)
    await file.seek(0)

    if file_type == "unknown":
        raise HTTPException(status_code=400, detail=ERROR_MESSAGES["unsupported_file"])

    file_content = await file.read()
    file_hash = hashlib.md5(file_content).hexdigest()
    file_identifier = f"upload_{file_hash}"

    return await process_queries(
        file_identifier,
        questions,
        is_url=False,
        file_content=file_content,
        filename=file.filename,
    )


@app.post("/hackrx/run/simple", response_model=SimpleQueryResponse)
async def run_query_simple(request: QueryRequest):
    full_response = await process_queries(
        request.documents, request.questions, is_url=True
    )
    return SimpleQueryResponse(
        answers=[answer.answer for answer in full_response.answers]
    )


async def process_queries(
    document_identifier: str,
    questions: List[str],
    is_url: bool = True,
    file_content: bytes = None,
    filename: str = None,
) -> QueryResponse:
    start_time = time.time()

    try:
        # get document content
        if is_url:
            file_content = await download_pdf_cached(document_identifier)
            doc_id = hashlib.md5(document_identifier.encode()).hexdigest()
            filename = None
        else:
            doc_id = document_identifier

        # extract text
        documents = extract_structured_text_from_document(
            file_content, doc_id, filename
        )
        if not documents:
            raise HTTPException(
                status_code=400, detail=ERROR_MESSAGES["no_text_extracted"]
            )

        # setup retriever
        retriever = await setup_hybrid_retriever_cached(documents, doc_id)

        # process questions concurrently
        async def process_single_question(question: str) -> Answer:
            try:
                relevant_chunks = await get_relevant_context_with_reranking(
                    question, retriever
                )
                result = await generate_structured_answer(question, relevant_chunks)

                return Answer(
                    query=question,
                    answer=result["answer"],
                    rationale=result["rationale"],
                    source=result["source"],
                )
            except Exception as e:
                return Answer(
                    query=question,
                    answer=f"processing error: {str(e)}",
                    rationale="processing failed",
                    source="error",
                )

        # controlled concurrency for speed
        semaphore = asyncio.Semaphore(CONCURRENT_QUESTION_LIMIT)

        async def bounded_process_question(question: str) -> Answer:
            async with semaphore:
                return await process_single_question(question)

        tasks = [bounded_process_question(q) for q in questions]
        answers = await asyncio.gather(*tasks)

        processing_time = time.time() - start_time
        logger.info(f"processed {len(questions)} questions in {processing_time:.2f}s")

        return QueryResponse(answers=answers)

    except Exception as e:
        logger.error(f"pipeline error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"processing failed: {str(e)}")


@app.get("/")
async def root():
    return {
        "message": "hackrx query system",
        "status": "ready",
        "supported_formats": ["pdf", "docx", "email"],
        "gpu_available": torch.cuda.is_available(),
    }


@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "gpu_available": torch.cuda.is_available(),
        "cached_documents": len(DOCUMENT_CACHE),
        "cached_retrievers": len(RETRIEVER_CACHE),
        "gemini_configured": gemini_model is not None,
    }


@app.on_event("startup")
async def startup_event():
    try:
        initialize_components()
        logger.info("startup completed")
    except Exception as e:
        logger.error(f"startup failed: {str(e)}")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")
